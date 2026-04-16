"""Office 文档加载器 - 支持 Word、Excel"""

from pathlib import Path
from typing import List

from llama_index.core.schema import Document

from src.rag.components.loaders.base import BaseLoader


class DocxLoader(BaseLoader):
    """Word 文档加载器 (.docx)"""

    def load_data(self, file_path: Path) -> List[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables_md = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.replace('|', '\\|') for cell in row.cells]
                rows.append('| ' + ' | '.join(cells) + ' |')
            if rows and len(table.rows) > 0:
                col_count = len(table.rows[0].cells)
                rows.insert(1, '|' + '---|' * col_count)
            tables_md.append('\n'.join(rows))

        content = '\n\n'.join(paragraphs)
        if tables_md:
            content += '\n\n## 表格\n\n' + '\n\n'.join(tables_md)

        metadata = {
            'source': str(file_path),
            'type': 'docx',
            'paragraphs': len(paragraphs),
            'tables': len(doc.tables),
        }

        return [Document(text=content, metadata=metadata)]


class ExcelLoader(BaseLoader):
    """Excel 加载器 (.xlsx, .xls)"""

    def __init__(self, max_rows: int = 10000):
        self.max_rows = max_rows

    def load_data(self, file_path: Path) -> List[Document]:
        import pandas as pd

        xl_file = pd.ExcelFile(file_path)
        documents = []

        for sheet_name in xl_file.sheet_names:
            df = xl_file.parse(sheet_name, nrows=self.max_rows)
            markdown = df.to_markdown(index=False)
            content = f'## Sheet: {sheet_name}\n\n{markdown}'

            if len(df) >= self.max_rows:
                content += f'\n\n> ⚠️ 表格过大，仅显示前 {self.max_rows} 行'

            metadata = {
                'source': str(file_path),
                'type': 'excel',
                'sheet': sheet_name,
                'rows': len(df),
                'columns': list(df.columns),
            }
            documents.append(Document(text=content, metadata=metadata))

        return documents


class UnifiedOfficeLoader(BaseLoader):
    """统一 Office 加载器 - 自动识别文件类型"""

    LOADERS = {
        '.docx': DocxLoader,
        '.xlsx': ExcelLoader,
        '.xls': ExcelLoader,
    }

    def load_data(self, file_path: Path) -> List[Document]:
        suffix = Path(file_path).suffix.lower()
        loader_class = self.LOADERS.get(suffix)

        if not loader_class:
            raise ValueError(f'不支持的文件格式: {suffix}')

        return loader_class().load_data(file_path)
